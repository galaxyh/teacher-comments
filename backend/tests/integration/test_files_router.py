"""End-to-end test for /file/{id}/process via TestClient.

Verifies the full HTTP layer: session cookie → router → ProcessingPipeline →
processed_artifact written → response.
"""

from __future__ import annotations

import io
import subprocess
from decimal import Decimal

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.adapters.openrouter_client import ChatResult


class FakeDriveClient:
    def __init__(self, *, file_bytes: bytes) -> None:
        self.file_bytes = file_bytes

    async def download_file(self, *, drive_file_id: str) -> bytes:
        return self.file_bytes


class FakeOpenRouter:
    def __init__(self, *, response_text: str = "## 摘要\n- 學生表現積極") -> None:
        self.response_text = response_text

    async def chat(self, **kwargs):
        return ChatResult(
            text=self.response_text,
            model_used=kwargs["model_id"],
            input_tokens=100,
            output_tokens=50,
            cost_usd=Decimal("0.0001"),
        )


@pytest.fixture
def authed_client(isolated_env):
    """TestClient with a real session cookie + fakes injected via dependency overrides."""
    subprocess.run(
        ["uv", "run", "alembic", "upgrade", "head"],
        check=True,
        cwd="/home/steven/projects/teacher-comments/backend",
        capture_output=True,
    )

    from app.adapters.document_extractors import build_default_registry
    from app.config import get_settings
    from app.core.session import issue_session_cookie, COOKIE_NAME
    from app.db.write_queue import get_write_queue
    from app.main import create_app
    from app.models import DriveFile, Teacher
    from app.routers.files import get_processing_pipeline
    from app.services.audit_logger import AuditLogger
    from app.services.llm_service import LLMService
    from app.services.pii_anonymizer import PIIAnonymizer
    from app.services.processing_pipeline import ProcessingPipeline

    # Build a minimal .docx
    doc = Document()
    doc.add_heading("週記", level=1)
    doc.add_paragraph("學生今天表現很好。0912345678 是家長電話。")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    fake_drive = FakeDriveClient(file_bytes=docx_bytes)
    fake_or = FakeOpenRouter()

    app = create_app()

    # Use the app's queue (started via lifespan inside TestClient context)
    def _override_pipeline() -> ProcessingPipeline:
        queue = get_write_queue()
        audit = AuditLogger(queue)
        anonymizer = PIIAnonymizer(db_write_queue=queue)
        llm = LLMService(
            settings=get_settings(),
            anonymizer=anonymizer,
            openrouter=fake_or,  # type: ignore[arg-type]
            audit=audit,
            db_write_queue=queue,
        )

        async def factory(*, teacher_id: str):
            return fake_drive

        return ProcessingPipeline(
            drive_client_factory=factory,  # type: ignore[arg-type]
            extractors=build_default_registry(),
            llm=llm,
        )

    app.dependency_overrides[get_processing_pipeline] = _override_pipeline

    with TestClient(app) as client:
        # Seed teacher + drive_file directly via the started write queue
        import asyncio

        queue = get_write_queue()

        async def seed_teacher(session):
            session.add(Teacher(id="t1", google_sub="sub", email="t@e.com"))

        async def seed_drive_file(session):
            session.add(
                DriveFile(
                    id="df1",
                    teacher_id="t1",
                    drive_file_id="drive-id-1",
                    semester_label="113-1",
                    student_pseudo_id="王小明",
                    category="learning",
                    drive_path="...",
                    filename="週記.docx",
                    mime_type=(
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ),
                    drive_modified_at="2026-05-01T00:00:00Z",
                )
            )

        # Submit synchronously by running a temporary loop via TestClient's portal
        client.portal.call(queue.submit, seed_teacher)
        client.portal.call(queue.submit, seed_drive_file)

        # Issue a session cookie for t1
        client.cookies.set(COOKIE_NAME, issue_session_cookie("t1"))
        yield client


def test_process_file_happy_path(authed_client: TestClient) -> None:
    r = authed_client.post("/file/df1/process")
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["artifact"]["state"] == "processed"
    assert body["artifact"]["llm_tier"] == "summary_cheap"
    assert "摘要" in body["artifact"]["content_markdown"]
    assert body["cost_usd"] > 0


def test_process_file_404_for_unknown_id(authed_client: TestClient) -> None:
    r = authed_client.post("/file/does-not-exist/process")
    assert r.status_code == 404
    assert r.json()["detail"]["reason"] == "drive_file_not_found"


def test_process_file_anonymous_returns_401(authed_client: TestClient) -> None:
    authed_client.cookies.clear()
    r = authed_client.post("/file/df1/process")
    assert r.status_code == 401


def test_get_artifact_after_processing(authed_client: TestClient) -> None:
    # Process first
    r1 = authed_client.post("/file/df1/process")
    assert r1.status_code == 200

    # Then GET the artifact
    r2 = authed_client.get("/file/df1/artifact")
    assert r2.status_code == 200
    body = r2.json()
    assert body["state"] == "processed"
    assert body["content_markdown"] is not None


def test_get_artifact_404_when_no_processing_done(authed_client: TestClient) -> None:
    r = authed_client.get("/file/df1/artifact")
    assert r.status_code == 404
    assert r.json()["detail"]["reason"] == "no_artifact_yet"
