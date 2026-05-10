"""ProcessingPipeline e2e: download → extract → LLM → restore → return ProcessingResult.

Uses fake DriveClient (canned bytes) + fake OpenRouter (canned response) + real
PIIAnonymizer + real DocumentExtractorRegistry. Verifies the full chokepoint
chain works against the real DB.
"""

from __future__ import annotations

import io
import subprocess
from decimal import Decimal

import pytest
from docx import Document

from app.adapters.document_extractors import build_default_registry
from app.adapters.openrouter_client import ChatResult


class FakeDriveClient:
    """Minimal stub — only download_file is exercised in pipeline tests."""

    def __init__(self, *, file_bytes: bytes) -> None:
        self.file_bytes = file_bytes
        self.download_calls: list[str] = []

    async def download_file(self, *, drive_file_id: str) -> bytes:
        self.download_calls.append(drive_file_id)
        return self.file_bytes


class FakeOpenRouter:
    def __init__(self, *, response_text: str = "## 摘要\n- 觀察 1\n- 觀察 2") -> None:
        self.response_text = response_text
        self.last_prompt: str | None = None

    async def chat(
        self,
        *,
        model_id: str,
        prompt: str,
        max_output_tokens: int = 1024,
        temperature: float = 0.7,
        timeout: float = 60.0,
    ) -> ChatResult:
        self.last_prompt = prompt
        return ChatResult(
            text=self.response_text,
            model_used=model_id,
            input_tokens=200,
            output_tokens=80,
            cost_usd=Decimal("0.00002"),
        )


@pytest.fixture
async def pipeline_harness(isolated_env, write_queue):
    """Real PIIAnonymizer + LLMService chokepoint, fake Drive + OpenRouter."""
    subprocess.run(
        ["uv", "run", "alembic", "upgrade", "head"],
        check=True,
        cwd="/home/steven/projects/teacher-comments/backend",
        capture_output=True,
    )

    from app.config import get_settings
    from app.models import DriveFile, Teacher
    from app.services.audit_logger import AuditLogger
    from app.services.llm_service import LLMService
    from app.services.pii_anonymizer import PIIAnonymizer
    from app.services.processing_pipeline import ProcessingPipeline

    # Build a real .docx in memory
    doc = Document()
    doc.add_heading("週記第三週", level=1)
    doc.add_paragraph("今天王小明完成了實驗。聯絡電話 0912345678。")
    doc.add_paragraph("觀察：學生表現積極。")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    # Seed teacher first (separate commit so FK is satisfied before DriveFile insert)
    async def insert_teacher(session) -> str:
        session.add(Teacher(id="t1", google_sub="sub", email="t@e.com"))
        return "t1"
    await write_queue.submit(insert_teacher)

    async def insert_drive_file(session) -> str:
        f = DriveFile(
            id="df1",
            teacher_id="t1",
            drive_file_id="drive-id-1",
            semester_label="113-1",
            student_pseudo_id="王小明",
            category="learning",
            drive_path="113-1/王小明/學習紀錄/週記.docx",
            filename="週記.docx",
            mime_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            drive_modified_at="2026-05-01T00:00:00Z",
        )
        session.add(f)
        return "df1"
    await write_queue.submit(insert_drive_file)

    fake_drive = FakeDriveClient(file_bytes=docx_bytes)
    fake_or = FakeOpenRouter()
    audit = AuditLogger(write_queue)
    anonymizer = PIIAnonymizer(db_write_queue=write_queue)
    llm = LLMService(
        settings=get_settings(),
        anonymizer=anonymizer,
        openrouter=fake_or,  # type: ignore[arg-type]
        audit=audit,
        db_write_queue=write_queue,
    )

    async def factory(*, teacher_id: str):
        return fake_drive

    pipeline = ProcessingPipeline(
        drive_client_factory=factory,  # type: ignore[arg-type]
        extractors=build_default_registry(),
        llm=llm,
    )

    # Look up the drive_file
    from app.db.session import get_sessionmaker
    async with get_sessionmaker()() as s:
        df = await s.get(DriveFile, "df1")

    return pipeline, df, fake_drive, fake_or, write_queue


@pytest.mark.asyncio
async def test_pipeline_processes_docx_end_to_end(pipeline_harness) -> None:
    pipeline, drive_file, fake_drive, fake_or, _queue = pipeline_harness

    result = await pipeline.process(teacher_id="t1", drive_file=drive_file)

    # Drive download invoked
    assert fake_drive.download_calls == ["drive-id-1"]
    # OpenRouter received an anonymised prompt — phone replaced with PH001
    assert "0912345678" not in fake_or.last_prompt
    assert "PH001" in fake_or.last_prompt
    # Pipeline output: PII-restored markdown summary + cost + content_hash
    assert result.artifact_type == "markdown_summary"
    assert result.content_markdown.startswith("## 摘要")
    assert result.llm_cost_usd == Decimal("0.00002")
    assert len(result.content_hash) == 64  # SHA-256 hex


@pytest.mark.asyncio
async def test_pipeline_unsupported_format_raises_terminal(pipeline_harness) -> None:
    """A drive_file row with a mime_type we don't route should raise UnsupportedFormatError."""
    from app.core.exceptions import UnsupportedFormatError
    from app.models import DriveFile

    pipeline, _, _, _, queue = pipeline_harness

    async def insert(session) -> str:
        f = DriveFile(
            id="df2",
            teacher_id="t1",
            drive_file_id="drive-id-2",
            semester_label="113-1",
            student_pseudo_id="王小明",
            category="work",
            drive_path="...",
            filename="image.png",
            mime_type="image/png",
            drive_modified_at="2026-05-01T00:00:00Z",
        )
        session.add(f)
        return "df2"
    await queue.submit(insert)

    from app.db.session import get_sessionmaker
    async with get_sessionmaker()() as s:
        unsupported = await s.get(DriveFile, "df2")

    with pytest.raises(UnsupportedFormatError):
        await pipeline.process(teacher_id="t1", drive_file=unsupported)
