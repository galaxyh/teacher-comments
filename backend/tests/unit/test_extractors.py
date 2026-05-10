"""Document extractor tests — DocxExtractor + PlainTextExtractor + registry."""

from __future__ import annotations

import io

import pytest
from docx import Document

from app.adapters.document_extractors import (
    DocumentExtractorRegistry,
    build_default_registry,
)
from app.adapters.document_extractors.docx import (
    _OLE_ENCRYPTED_SIG,
    _OLE_HEADER_MAGIC,
    DocxExtractor,
)
from app.adapters.document_extractors.text import PlainTextExtractor
from app.core.exceptions import UnsupportedFormatError


def _build_docx(*, paragraphs: list[tuple[str, str]] | None = None) -> bytes:
    """Build an in-memory .docx. paragraphs = [(text, style_name)]."""
    doc = Document()
    for text, style in (paragraphs or [("Hello world", "Normal")]):
        p = doc.add_paragraph(text)
        if style and style != "Normal":
            try:
                p.style = doc.styles[style]
            except KeyError:
                pass  # style not in default template; tests can still proceed
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxExtractor:
    @pytest.mark.asyncio
    async def test_plain_paragraphs(self) -> None:
        ext = DocxExtractor()
        body = _build_docx(paragraphs=[("Hello", "Normal"), ("World", "Normal")])
        result = await ext.extract(file_bytes=body, filename="hello.docx")
        # Two paragraphs joined by blank-line separator
        assert "Hello" in result.text and "World" in result.text
        assert result.has_images is False

    @pytest.mark.asyncio
    async def test_heading_styles_become_markdown(self) -> None:
        ext = DocxExtractor()
        body = _build_docx(
            paragraphs=[("Title", "Heading 1"), ("Sub", "Heading 2"), ("Body", "Normal")]
        )
        result = await ext.extract(file_bytes=body, filename="t.docx")
        assert "# Title" in result.text
        assert "## Sub" in result.text

    @pytest.mark.asyncio
    async def test_table_renders_as_pipe_table(self) -> None:
        ext = DocxExtractor()
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "h1"
        table.cell(0, 1).text = "h2"
        table.cell(1, 0).text = "v1"
        table.cell(1, 1).text = "v2"
        buf = io.BytesIO()
        doc.save(buf)

        result = await ext.extract(file_bytes=buf.getvalue(), filename="t.docx")
        # GFM-style header / separator / body
        assert "| h1 | h2 |" in result.text
        assert "| --- | --- |" in result.text
        assert "| v1 | v2 |" in result.text

    @pytest.mark.asyncio
    async def test_encrypted_ole_raises_unsupported(self) -> None:
        ext = DocxExtractor()
        # Synthesize an OLE header + encrypted signature in first 8KiB
        body = _OLE_HEADER_MAGIC + b"\x00" * 100 + _OLE_ENCRYPTED_SIG + b"\x00" * 1000
        with pytest.raises(UnsupportedFormatError, match="Encrypted"):
            await ext.extract(file_bytes=body, filename="locked.docx")

    @pytest.mark.asyncio
    async def test_legacy_doc_raises_unsupported(self) -> None:
        ext = DocxExtractor()
        body = _OLE_HEADER_MAGIC + b"\x00" * 1000  # OLE but not encrypted = legacy .doc
        with pytest.raises(UnsupportedFormatError, match="Legacy"):
            await ext.extract(file_bytes=body, filename="old.docx")

    def test_supports_via_mime_or_filename(self) -> None:
        ext = DocxExtractor()
        assert ext.supports(
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="x",
        )
        assert ext.supports(mime_type="application/octet-stream", filename="x.docx")
        assert not ext.supports(mime_type="text/plain", filename="x.txt")


class TestPlainTextExtractor:
    @pytest.mark.asyncio
    async def test_utf8_happy_path(self) -> None:
        ext = PlainTextExtractor()
        body = "你好 hello".encode("utf-8")
        result = await ext.extract(file_bytes=body, filename="x.txt")
        assert result.text == "你好 hello"
        assert result.warnings == []

    @pytest.mark.asyncio
    async def test_big5_fallback(self) -> None:
        ext = PlainTextExtractor()
        body = "舊式繁中".encode("big5")
        result = await ext.extract(file_bytes=body, filename="legacy.txt")
        assert result.text == "舊式繁中"
        assert any("big5" in w for w in result.warnings)

    @pytest.mark.asyncio
    async def test_undecodable_uses_replace(self) -> None:
        ext = PlainTextExtractor()
        # Bytes that are valid in none of utf-8 / big5 / cp950 cleanly — use a synthesized
        # mid-multibyte truncation
        body = b"\xff\xfe\xfd\xfc"
        result = await ext.extract(file_bytes=body, filename="garbage.txt")
        # Some output produced; warning recorded
        assert any("replaced" in w or "unable" in w for w in result.warnings)


class TestRegistry:
    def test_get_routes_to_first_supporting_extractor(self) -> None:
        reg = build_default_registry()
        # docx mime
        assert isinstance(
            reg.get(mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="x"),
            DocxExtractor,
        )
        # txt
        assert isinstance(
            reg.get(mime_type="text/plain", filename="x.txt"),
            PlainTextExtractor,
        )

    def test_get_raises_when_no_match(self) -> None:
        reg = DocumentExtractorRegistry(extractors=[])
        with pytest.raises(UnsupportedFormatError):
            reg.get(mime_type="application/x-bogus", filename="x.bogus")
